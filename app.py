import streamlit as st
import json
import os
from datetime import datetime
from io import BytesIO
from docx import Document
import unicodedata
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

# Configuração da porta
os.environ["STREAMLIT_SERVER_PORT"] = "8502"

# --- Configurações Iniciais ---
if getattr(sys, 'frozen', False):
    base_path = sys._MEIPASS
else:
    base_path = os.path.abspath(".")

DATA_DIR = os.path.join(base_path, "data")
USERS_FILE = os.path.join(DATA_DIR, "users.json")
PARECERES_FILE = os.path.join(DATA_DIR, "pareceres.json")
SCHOOL_NAME = "ESCOLA MUNICIPAL DE EDUCAÇÃO FUNDAMENTAL ELESBÃO BARBOSA DE CARVALHO"
COORDENADOR_NAME = "NOME DO COORDENADOR AQUI"

# --- Funções de Utilitário ---
def load_data(file_path, default_value):
    """Carrega dados de um arquivo JSON, retornando um valor padrão se não existir."""
    if not os.path.exists(file_path):
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(default_value, f, ensure_ascii=False, indent=4)
    with open(file_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def save_data(data, file_path):
    """Salva dados em um arquivo JSON."""
    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

def sanitize_student_name_for_filename(name):
    """Remove caracteres especiais e espaços do nome para usar em nomes de arquivo."""
    sanitized = ''.join(c for c in name if c.isalnum() or c.isspace())
    return sanitized.strip().replace(' ', '_')

def create_parecer_docx(data):
    """Cria um documento Word com o parecer do aluno a partir de um template."""
    template_path = os.path.join(DATA_DIR, "parecer_template.docx")
    try:
        doc = Document(template_path)
    except FileNotFoundError:
        return st.error(f"Erro: Arquivo de template não encontrado em {template_path}. Certifique-se de que o arquivo 'parecer_template.docx' está na pasta 'data'.")

    replacements = {
        'Nº ': data['numero_aluno'],
        'Período: ': data['periodo'],
        'Turma: ': data['turma'],
        'Turno: ': data['turno'],
        'Ano Letivo: ': data['ano_letivo'],
        'Semestre: ': data['semestre'],
        'Nome do Aluno (a):': data['nome_aluno'],
        'Filiação: ': f"Filiação: {data['filiacao_mae']} e {data['filiacao_pai']}",
        'Endereço: ': data['endereco'],
        'UF: ': data['uf'],
        'Data de Nascimento: ': data['data_nascimento'],
        'Naturalidade: ': data['naturalidade'],
        'Professor': data['nome_professor'],
        'Coordenador': COORDENADOR_NAME,
        'Maravilha, AL': f"Maravilha, AL, {data['data_parecer']}",
    }

    # Substituir no documento inteiro
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)

    # Inserir o corpo do parecer
    for paragraph in doc.paragraphs:
        if 'PARECER DESCRITIVO:' in paragraph.text:
            p = paragraph.insert_paragraph_before(data['texto_parecer'])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            break

    # Salvar em BytesIO para download
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- Funções de Autenticação ---
def login():
    st.session_state['logged_in'] = False
    
    # Interface melhorada para login
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("## 📝 Sistema de Parecer Descritivo")
        st.markdown(f"### {SCHOOL_NAME}")
        st.markdown("---")
        
        with st.form("login_form"):
            st.markdown("#### Acesso ao Sistema")
            username = st.text_input("👤 Usuário", placeholder="Digite seu usuário")
            password = st.text_input("🔒 Senha", type="password", placeholder="Digite sua senha")
            
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                login_btn = st.form_submit_button("🚀 Entrar", use_container_width=True)
            with col_btn2:
                create_btn = st.form_submit_button("➕ Criar Conta", use_container_width=True)
            
            if login_btn:
                users = load_data(USERS_FILE, {})
                if username in users and users[username]['password'] == password:
                    st.session_state['logged_in'] = True
                    st.session_state['username'] = username
                    st.success("✅ Login bem-sucedido!")
                    st.rerun()
                else:
                    st.error("❌ Usuário ou senha incorretos.")
            
            if create_btn:
                st.session_state['show_create_account'] = True
                st.rerun()

def create_account():
    st.session_state['logged_in'] = False
    st.session_state['show_create_account'] = True
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("## 📝 Sistema de Parecer Descritivo")
        st.markdown("### Criar Nova Conta")
        st.markdown("---")
        
        with st.form("create_account_form"):
            new_username = st.text_input("👤 Novo Usuário", key="new_user", placeholder="Escolha um nome de usuário")
            new_password = st.text_input("🔒 Nova Senha", type="password", key="new_pass", placeholder="Crie uma senha segura")
            confirm_password = st.text_input("🔒 Confirmar Senha", type="password", key="confirm_pass", placeholder="Confirme sua senha")
            
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                confirm_btn = st.form_submit_button("✅ Confirmar Cadastro", use_container_width=True)
            with col_btn2:
                back_btn = st.form_submit_button("⬅️ Voltar", use_container_width=True)
            
            if confirm_btn:
                if not new_username or not new_password or not confirm_password:
                    st.error("⚠️ Preencha todos os campos.")
                elif new_password != confirm_password:
                    st.error("⚠️ As senhas não coincidem.")
                else:
                    users = load_data(USERS_FILE, {})
                    if new_username in users:
                        st.error("⚠️ Usuário já existe. Escolha outro nome.")
                    else:
                        users[new_username] = {'password': new_password, 'is_admin': False}
                        save_data(users, USERS_FILE)
                        st.success("✅ Conta criada com sucesso! Faça login para continuar.")
                        st.session_state['show_create_account'] = False
                        st.rerun()
            
            if back_btn:
                st.session_state['show_create_account'] = False
                st.rerun()

def generate_parecer_text(student_name, status_aluno, foco_parecer, **kwargs):
    """Gera o texto do parecer baseado nos parâmetros fornecidos."""
    
    if status_aluno == "Deixou de Frequentar":
        return (
            f"Constatou-se que o(a) aluno(a) {student_name} deixou de frequentar a escola durante "
            "o período letivo. Devido ao curto período de tempo de sua presença em sala de aula, "
            "não foi possível estabelecer uma relação de aprendizado sólida, bem como avaliar "
            "adequadamente seu desenvolvimento nas diferentes áreas do conhecimento. A ausência "
            "prolongada impossibilitou a construção de vínculos pedagógicos consistentes e o "
            "acompanhamento do processo de ensino-aprendizagem."
        )
    
    elif status_aluno == "Transferido":
        return (
            f"O(a) aluno(a) {student_name} foi transferido(a) durante o período letivo. "
            "Durante sua permanência na instituição, demonstrou estar em processo de adaptação "
            "ao ambiente escolar. O tempo de convivência foi insuficiente para uma avaliação "
            "completa de seu desenvolvimento acadêmico e social, mas observou-se potencial "
            "para o aprendizado nas atividades propostas."
        )
    
    elif status_aluno == "Necessidades Especiais":
        return (
            f"O(a) aluno(a) {student_name} apresenta necessidades educacionais especiais e "
            "recebe acompanhamento pedagógico diferenciado. Demonstra progresso gradual em "
            "seu desenvolvimento, respeitando-se suas particularidades e ritmo de aprendizagem. "
            "As atividades são adaptadas às suas necessidades, promovendo sua inclusão e "
            "participação efetiva no processo educativo."
        )
    
    elif status_aluno == "Frequente":
        if foco_parecer == "Geral":
            leitura = kwargs.get('leitura', 'N/A')
            matematica = kwargs.get('matematica', 'N/A')
            comportamento = kwargs.get('comportamento', 'N/A')
            participacao = kwargs.get('participacao', 'N/A')
            
            return (
                f"O(a) aluno(a) {student_name} demonstra uma postura colaborativa em sala de aula, "
                f"apresentando um desempenho geral satisfatório. No que diz respeito à leitura e "
                f"escrita, seu desenvolvimento é {leitura.lower()}. Na área de Matemática, "
                f"demonstra desempenho {matematica.lower()}. Quanto ao comportamento, apresenta "
                f"conduta {comportamento.lower()} e sua participação nas atividades propostas "
                f"é {participacao.lower()}, contribuindo positivamente para o ambiente de aprendizagem."
            )
        
        elif foco_parecer == "Comportamental":
            comportamento = kwargs.get('comportamento', 'N/A')
            participacao = kwargs.get('participacao', 'N/A')
            
            return (
                f"O(a) aluno(a) {student_name} demonstra uma postura {comportamento.lower()} "
                f"em sala de aula, mantendo relacionamento respeitoso com colegas e professores. "
                f"Sua participação nas atividades é {participacao.lower()}, mostrando interesse "
                "pelas propostas pedagógicas e contribuindo de forma positiva para o ambiente escolar."
            )
        
        elif foco_parecer == "Específico":
            return kwargs.get('parecer_personalizado', '')
    
    return ""

# --- Conteúdo do Aplicativo ---
def app_content():
    # Sidebar melhorada
    with st.sidebar:
        st.markdown(f"## 👋 Bem-vindo(a)")
        st.markdown(f"**{st.session_state['username']}**")
        st.markdown("---")
        
        if st.button("🚪 Logout", use_container_width=True):
            st.session_state.clear()
            st.rerun()
        
        st.markdown("---")
        st.markdown("### 📊 Estatísticas")
        pareceres_db = load_data(PARECERES_FILE, {})
        user_pareceres = pareceres_db.get(st.session_state['username'], {})
        total_alunos = len(user_pareceres)
        total_pareceres = sum(len(pareceres) for pareceres in user_pareceres.values())
        
        st.metric("Total de Alunos", total_alunos)
        st.metric("Total de Pareceres", total_pareceres)

    # Título principal
    st.markdown("# 📝 Gerador de Parecer Descritivo Individual")
    st.markdown(f"### {SCHOOL_NAME}")
    st.markdown("---")
    
    # Seletor de tipo de parecer
    st.markdown("## 📄 Tipo de Documento")
    parecer_type = st.radio(
        "Selecione o tipo de parecer a ser gerado:",
        ["📋 Parecer Completo", "📝 Apenas o Texto"],
        horizontal=True
    )

    # Dados do aluno
    st.markdown("## 👤 Dados do Aluno")
    with st.container():
        col1, col2 = st.columns(2)
        with col1:
            student_name = st.text_input("📝 Nome Completo do Aluno(a)", placeholder="Digite o nome completo")
            filiacao_mae = st.text_input("👩 Filiação (Mãe)", placeholder="Nome da mãe")
            data_nascimento = st.text_input("📅 Data de Nascimento", placeholder="dd/mm/aaaa")
            endereco = st.text_input("🏠 Endereço", placeholder="Endereço completo")
            numero_aluno = st.text_input("📋 Nº na Lista de Chamada", placeholder="Número do aluno")
        
        with col2:
            filiacao_pai = st.text_input("👨 Filiação (Pai)", placeholder="Nome do pai")
            naturalidade = st.text_input("🌍 Naturalidade", placeholder="Cidade de nascimento")
            uf = st.text_input("📍 UF", placeholder="Estado", max_chars=2)
            # Espaços em branco para alinhamento
            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown("<br>", unsafe_allow_html=True)

    # Dados da turma
    st.markdown("## 🏫 Dados da Turma")
    with st.container():
        col1, col2, col3 = st.columns(3)
        with col1:
            periodo = st.selectbox("📚 Período/Ano", 
                                 ["1º", "2º", "3º", "4º", "5º", "6º", "7º", "8º", "9º", "EJA"])
        with col2:
            turma = st.text_input("🏷️ Turma", placeholder="Ex: A, B, Única")
        with col3:
            turno = st.selectbox("⏰ Turno", ["Manhã", "Tarde", "Noturno"])

    # Dados do parecer
    st.markdown("## 👨‍🏫 Dados do Parecer")
    nome_professor = st.text_input("🧑‍🏫 Nome Completo do(a) Professor(a)", 
                                  placeholder="Digite o nome completo do professor")
    
    # Status do aluno (MELHORADO)
    st.markdown("## 📊 Status do Aluno")
    status_aluno = st.selectbox(
        "📋 Situação do Aluno(a)",
        [
            "Frequente", 
            "Deixou de Frequentar", 
            "Transferido",
            "Necessidades Especiais"
        ]
    )

    # Interface condicional baseada no status
    if status_aluno == "Frequente":
        st.markdown("### 🎯 Foco do Parecer")
        foco_parecer = st.radio(
            "Selecione o foco principal:", 
            ["📊 Geral", "🤝 Comportamental", "✏️ Específico"],
            horizontal=True
        )
    
        if foco_parecer == "📊 Geral":
            st.markdown("#### 📈 Áreas de Desempenho")
            with st.container():
                col1, col2 = st.columns(2)
                with col1:
                    leitura = st.selectbox("📖 Leitura e Escrita", 
                                         ["Excelente", "Muito Bom", "Bom", "Satisfatório", "Precisa Melhorar"])
                    comportamento = st.selectbox("😊 Comportamento", 
                                                ["Excelente", "Muito Bom", "Bom", "Satisfatório", "Precisa Melhorar"])
                with col2:
                    matematica = st.selectbox("🔢 Matemática", 
                                            ["Excelente", "Muito Bom", "Bom", "Satisfatório", "Precisa Melhorar"])
                    participacao = st.selectbox("🙋 Participação", 
                                               ["Excelente", "Muito Bom", "Bom", "Satisfatório", "Precisa Melhorar"])
        
        elif foco_parecer == "🤝 Comportamental":
            st.markdown("#### 🎭 Aspectos Comportamentais")
            with st.container():
                col1, col2 = st.columns(2)
                with col1:
                    comportamento = st.selectbox("😊 Comportamento", 
                                                ["Excelente", "Muito Bom", "Bom", "Satisfatório", "Precisa Melhorar"])
                with col2:
                    participacao = st.selectbox("🙋 Participação", 
                                               ["Excelente", "Muito Bom", "Bom", "Satisfatório", "Precisa Melhorar"])
                leitura = matematica = "N/A"
        
        elif foco_parecer == "✏️ Específico":
            st.markdown("#### 📝 Parecer Personalizado")
            parecer_personalizado = st.text_area(
                "Digite o texto do parecer:",
                placeholder="Descreva de forma detalhada o desempenho e desenvolvimento do aluno...",
                height=150
            )
            leitura = matematica = comportamento = participacao = "N/A"
    else:
        # Para alunos não frequentes, definir variáveis padrão
        foco_parecer = None
        leitura = matematica = comportamento = participacao = "N/A"
        parecer_personalizado = ""

    # Botão de geração (MELHORADO)
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        gerar_parecer = st.button("🚀 Gerar Parecer", use_container_width=True, type="primary")

    # Lógica de geração do parecer
    if gerar_parecer:
        if not student_name or not nome_professor:
            st.error("⚠️ Por favor, preencha o nome do aluno e do professor.")
        else:
            # Gerar texto do parecer
            parecer_kwargs = {
                'leitura': leitura,
                'matematica': matematica,
                'comportamento': comportamento,
                'participacao': participacao
            }
            
            if foco_parecer == "✏️ Específico":
                parecer_kwargs['parecer_personalizado'] = parecer_personalizado
            
            parecer_text = generate_parecer_text(
                student_name, 
                status_aluno, 
                foco_parecer.replace("📊 ", "").replace("🤝 ", "").replace("✏️ ", "") if foco_parecer else None,
                **parecer_kwargs
            )
            
            # Data formatada
            data_parecer = datetime.now().strftime("%d de %B de %Y").replace(
                "January", "janeiro"
            ).replace("February", "fevereiro").replace("March", "março").replace(
                "April", "abril"
            ).replace("May", "maio").replace("June", "junho").replace(
                "July", "julho"
            ).replace("August", "agosto").replace("September", "setembro").replace(
                "October", "outubro"
            ).replace("November", "novembro").replace("December", "dezembro")

            parecer_data = {
                'nome_aluno': student_name,
                'filiacao_mae': filiacao_mae,
                'filiacao_pai': filiacao_pai,
                'data_nascimento': data_nascimento,
                'naturalidade': naturalidade,
                'endereco': endereco,
                'uf': uf,
                'numero_aluno': numero_aluno,
                'periodo': periodo,
                'turma': turma,
                'turno': turno,
                'ano_letivo': datetime.now().year,
                'semestre': '2025.1',
                'nome_professor': nome_professor,
                'texto_parecer': parecer_text,
                'data_parecer': data_parecer,
                'status_aluno': status_aluno
            }

            if foco_parecer:
                parecer_data['foco_parecer'] = foco_parecer
                if foco_parecer not in ["✏️ Específico"] and status_aluno == "Frequente":
                    parecer_data['characteristics_levels'] = {
                        "leitura": leitura,
                        "matematica": matematica,
                        "comportamento": comportamento,
                        "participacao": participacao
                    }

            # Gerar arquivo DOCX se necessário
            docx_data_bytes = None
            if "📋 Parecer Completo" in parecer_type:
                doc_bytes_io = create_parecer_docx(parecer_data)
                if doc_bytes_io:
                    docx_data_bytes = doc_bytes_io.getvalue()
                    parecer_data['docx_data'] = docx_data_bytes.hex()
            
            # Exibir resultado
            st.success("✅ Parecer gerado com sucesso!")
            
            # Preview do parecer
            with st.container():
                st.markdown("### 📋 Pré-visualização do Parecer")
                st.markdown("---")
                st.markdown(f"**Aluno(a):** {student_name}")
                st.markdown(f"**Professor(a):** {nome_professor}")
                st.markdown(f"**Status:** {status_aluno}")
                st.markdown("**Parecer:**")
                st.info(parecer_data['texto_parecer'])
                
                # Botão de download
                if docx_data_bytes:
                    sanitized_name = sanitize_student_name_for_filename(student_name)
                    file_name = f"parecer_descritivo_{sanitized_name}.docx"
                    
                    col1, col2, col3 = st.columns([1, 1, 1])
                    with col2:
                        st.download_button(
                            label="📥 Baixar Parecer (DOCX)",
                            data=docx_data_bytes,
                            file_name=file_name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            type="secondary"
                        )

            # Salvar no banco de dados
            pareceres_db = load_data(PARECERES_FILE, {})
            if st.session_state['username'] not in pareceres_db:
                pareceres_db[st.session_state['username']] = {}
            if student_name not in pareceres_db[st.session_state['username']]:
                pareceres_db[st.session_state['username']][student_name] = []
            
            parecer_data_to_save = parecer_data.copy()
            if docx_data_bytes:
                parecer_data_to_save['docx_data'] = docx_data_bytes.hex()
            
            pareceres_db[st.session_state['username']][student_name].append(parecer_data_to_save)
            save_data(pareceres_db, PARECERES_FILE)

def admin_dashboard():
    """Painel do Administrador com interface melhorada."""
    with st.sidebar:
        st.markdown(f"## 👑 Administrador")
        st.markdown(f"**{st.session_state['username']}**")
        st.markdown("---")
        
        if st.button("🚪 Logout", use_container_width=True):
            st.session_state.clear()
            st.rerun()

    st.markdown("# 👑 Painel do Administrador")
    st.markdown("---")
    
    # Gerenciar usuários
    st.markdown("## 👥 Gerenciar Usuários")
    users_db = load_data(USERS_FILE, {})
    user_list = [user for user in users_db if user != st.session_state['username']]

    if user_list:
        st.markdown("### 📋 Usuários Existentes")
        for user in user_list:
            is_admin = users_db[user].get('is_admin', False)
            with st.container():
                col1, col2, col3, col4 = st.columns([2, 1, 1, 1])
                with col1:
                    admin_badge = "👑" if is_admin else "👤"
                    st.markdown(f"{admin_badge} **{user}**")
                with col2:
                    if not is_admin:
                        if st.button("👑 Tornar Admin", key=f"make_admin_{user}", use_container_width=True):
                            users_db[user]['is_admin'] = True
                            save_data(users_db, USERS_FILE)
                            st.success(f"✅ {user} agora é administrador.")
                            st.rerun()
                with col3:
                    if is_admin:
                        if st.button("👤 Remover Admin", key=f"remove_admin_{user}", use_container_width=True):
                            users_db[user]['is_admin'] = False
                            save_data(users_db, USERS_FILE)
                            st.success(f"✅ {user} não é mais administrador.")
                            st.rerun()
                with col4:
                    if st.button("🗑️ Remover", key=f"remove_user_{user}", use_container_width=True):
                        del users_db[user]
                        save_data(users_db, USERS_FILE)
                        st.success(f"✅ Usuário {user} removido.")
                        st.rerun()
                st.markdown("---")
    else:
        st.info("ℹ️ Nenhum outro usuário cadastrado.")

    # Gerenciar pareceres
    st.markdown("## 📝 Gerenciar Pareceres")
    pareceres_db = load_data(PARECERES_FILE, {})
    
    if pareceres_db:
        st.markdown("### 📊 Estatísticas Gerais")
        total_users = len(pareceres_db)
        total_students = sum(len(user_data) for user_data in pareceres_db.values())
        total_pareceres = sum(len(pareceres) for user_data in pareceres_db.values() 
                            for pareceres in user_data.values())
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("👥 Usuários", total_users)
        with col2:
            st.metric("👤 Alunos", total_students)
        with col3:
            st.metric("📝 Pareceres", total_pareceres)
        
        st.markdown("---")
        st.markdown("### 📋 Pareceres por Usuário")
        
        for username_parecer, user_data in pareceres_db.items():
            with st.expander(f"👤 Pareceres de {username_parecer} ({len(user_data)} alunos)"):
                for student_name_display, pareceres in user_data.items():
                    st.markdown(f"**👤 Aluno(a):** {student_name_display}")
                    
                    for i, parecer_info in enumerate(pareceres):
                        with st.container():
                            col1, col2 = st.columns([3, 1])
                            with col1:
                                st.markdown(f"**📝 Parecer {i+1}** - {parecer_info.get('data_parecer', 'Data Indisponível')}")
                                st.markdown(f"**👨‍🏫 Professor(a):** {parecer_info['nome_professor']}")
                                st.markdown(f"**📊 Status:** {parecer_info.get('status_aluno', 'N/A')}")
                                
                                # Preview do texto (limitado)
                                texto_preview = parecer_info['texto_parecer'][:200] + "..." if len(parecer_info['texto_parecer']) > 200 else parecer_info['texto_parecer']
                                st.markdown(f"**📄 Texto:** {texto_preview}")
                                
                                if 'characteristics_levels' in parecer_info and parecer_info['characteristics_levels']:
                                    levels = parecer_info['characteristics_levels']
                                    st.markdown(f"📖 **Leitura:** {levels.get('leitura', 'N/A')} | "
                                              f"🔢 **Matemática:** {levels.get('matematica', 'N/A')} | "
                                              f"😊 **Comportamento:** {levels.get('comportamento', 'N/A')} | "
                                              f"🙋 **Participação:** {levels.get('participacao', 'N/A')}")
                            
                            with col2:
                                if 'docx_data' in parecer_info and parecer_info['docx_data']:
                                    try:
                                        docx_data_bytes = bytes.fromhex(parecer_info['docx_data'])
                                        sanitized_name = sanitize_student_name_for_filename(student_name_display)
                                        file_name = f"parecer_{sanitized_name}_{i+1}.docx"
                                        st.download_button(
                                            label="📥 Download",
                                            data=docx_data_bytes,
                                            file_name=file_name,
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            key=f"admin_download_docx_{student_name_display}_{i}",
                                            use_container_width=True
                                        )
                                    except ValueError:
                                        st.error("❌ Erro nos dados do DOCX")
                                else:
                                    st.info("📄 Apenas texto")
                        
                        st.markdown("---")
    else:
        st.info("ℹ️ Nenhum parecer salvo ainda.")

    st.markdown("---")
    st.markdown("### ℹ️ Informações do Sistema")
    st.info("Sistema desenvolvido com Streamlit e Python para geração de pareceres descritivos educacionais.")

# --- Função Principal ---
def main():
    """Função principal do aplicativo."""
    # Configuração da página
    st.set_page_config(
        page_title="Sistema de Parecer Descritivo",
        page_icon="📝",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # CSS customizado para melhorar a aparência
    st.markdown("""
        <style>
        .main > div {
            padding-top: 2rem;
        }
        .stButton > button {
            border-radius: 10px;
            border: none;
            padding: 0.5rem 1rem;
            font-weight: 500;
        }
        .stSelectbox > div > div {
            border-radius: 10px;
        }
        .stTextInput > div > div {
            border-radius: 10px;
        }
        .stTextArea > div > div {
            border-radius: 10px;
        }
        .metric-container {
            background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
            padding: 1rem;
            border-radius: 10px;
            color: white;
        }
        </style>
    """, unsafe_allow_html=True)
    
    # Inicialização das variáveis de sessão
    if 'logged_in' not in st.session_state:
        st.session_state['logged_in'] = False
    if 'show_create_account' not in st.session_state:
        st.session_state['show_create_account'] = False

    # Lógica principal de navegação
    users_db = load_data(USERS_FILE, {})
    current_user = st.session_state.get('username')

    if st.session_state['logged_in']:
        if users_db.get(current_user, {}).get('is_admin', False):
            admin_dashboard()
        else:
            app_content()
    else:
        if st.session_state['show_create_account']:
            create_account()
        else:
            login()

# --- Execução do Aplicativo ---
if __name__ == "__main__":
    main()